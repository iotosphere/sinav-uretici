from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import json
import io
import re
import datetime

# Parsing Fonksiyonu
def parse_questions_from_ai_response(questions_data):
    """
    AI'dan gelen veriyi (JSON string veya Python listesi) parse eder.
    """
    # Eğer zaten liste ise direkt döndür
    if isinstance(questions_data, list):
        return questions_data
    
    # Eğer dict ise ve içinde 'questions' veya 'sorular' varsa onu döndür
    if isinstance(questions_data, dict):
        if 'questions' in questions_data:
            return questions_data['questions']
        if 'sorular' in questions_data:
            return questions_data['sorular']
        return [questions_data] # Tek bir obje olabilir

    # String ise JSON olarak parse etmeye çalış
    try:
        data = json.loads(questions_data)
        if isinstance(data, dict) and 'questions' in data:
            return data['questions']
        elif isinstance(data, list):
            return data
        elif isinstance(data, dict) and 'sorular' in data:
            return data['sorular']
    except json.JSONDecodeError as e:
        print(f"KRİTİK HATA: AI'dan gelen veri JSON formatında değil veya hatalı. Hata: {e}")
        return []
    
    return []

# DOCX Oluşturma Fonksiyonu
def build_docx_exam(metadata, questions_json):
    
    doc = Document()
    questions = parse_questions_from_ai_response(questions_json)

    # Eğitim-öğretim yılı (Türkiye: Eylül-Haziran)
    _now = datetime.datetime.now()
    _sy_start = _now.year if _now.month >= 9 else _now.year - 1
    school_year = f"{_sy_start}-{_sy_start + 1}"
    
    # Sayfa kenar boşluklarını ayarla
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # --- 1. KISIM: SINAV BAŞLIK TABLOSU ---
    
    header_table = doc.add_table(rows=4, cols=4)
    header_table.style = 'Table Grid'
    
    # Hücreleri doldurma
    header_table.cell(0, 0).text = metadata.get('sinif', '……') + ' SINIF'
    header_table.cell(0, 1).merge(header_table.cell(0, 3)).text = metadata.get('okul', 'OKUL ADI')

    header_table.cell(1, 0).merge(header_table.cell(1, 1)).text = f"{school_year} EĞİTİM ÖĞRETİM YILI"
    header_table.cell(1, 2).text = "ADI:"
    header_table.cell(1, 3).text = "PUAN"

    header_table.cell(2, 0).merge(header_table.cell(2, 1)).text = f"{metadata.get('ders', 'Ders Adı')} DERSİ"
    header_table.cell(2, 2).text = "SOYADI:"
    header_table.cell(2, 3).text = ""

    header_table.cell(3, 0).merge(header_table.cell(3, 1)).text = f"{metadata.get('sinav_no', '1. DÖNEM 1. YAZILI')} (……. SENARYO)"
    header_table.cell(3, 2).text = "SINIFI:"
    header_table.cell(3, 3).text = "OKUL NO:"
    
    # Tablonun ortalanması ve font ayarı
    for row in header_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph('\n')
    
    # --- 2. KISIM: PUAN TABLOSU (Toplam 100 Puan) ---
    
    question_count = len(questions) if questions else 10
    puan_per_soru = 100 // question_count
    kalan_puan = 100 - (puan_per_soru * question_count)
    
    # Dinamik kolon sayisi (max 10 soru)
    cols_needed = min(question_count, 10)
    score_table = doc.add_table(rows=2, cols=cols_needed * 2)
    score_table.style = 'Table Grid'
    
    for i in range(cols_needed):
        # Son sorulara kalan puanı ekle (eşit dağılım için)
        soru_puan = puan_per_soru + (1 if i >= (cols_needed - kalan_puan) else 0)
        
        q_cell = score_table.cell(0, i * 2)
        q_cell.merge(score_table.cell(0, i * 2 + 1))
        q_cell.text = f"Soru {i+1}\n({soru_puan} Puan)"
        q_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        score_table.cell(1, i * 2).merge(score_table.cell(1, i * 2 + 1)).text = "........."
        score_table.cell(1, i * 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Toplam puan bilgisi
    toplam_para = doc.add_paragraph()
    toplam_run = toplam_para.add_run(f"Toplam: 100 Puan")
    toplam_run.font.size = Pt(10)
    toplam_run.bold = True
    toplam_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('\n')
    
    # --- 3. KISIM: AÇIK UÇLU SORULAR VE CEVAP ALANLARI ---
    
    for i, item in enumerate(questions[:10], 1):
        # Item'i string olarak ele al
        if isinstance(item, str):
            soru_metni = item
            is_acik_uclu = True
        else:
            soru_metni = item.get('soru_metni', item.get('soru', f"Soru {i} metni okunamadı."))
            soru_tipi = item.get('soru_tipi', '')
            # Seçenek varsa test, yoksa açık uçlu
            is_acik_uclu = not item.get('secenekler') or len(item.get('secenekler', [])) == 0

        # Soru başlığı (sadece soru numarası - öğrenci görmesin)
        q_heading_para = doc.add_paragraph()
        soru_bilgisi = f"Soru {i})"
        
        # Not: Bloom, zorluk, kazanım bilgileri öğrenci kağıdında GÖSTERİLMİYOR
        # Bu bilgiler sadece cevap anahtarında (öğretmen için) yer alıyor
        
        q_heading_run = q_heading_para.add_run(soru_bilgisi)
        q_heading_run.font.size = Pt(10)
        q_heading_run.bold = True
        q_heading_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)  # Koyu gri
        
        # Soru metni
        q_text_para = doc.add_paragraph()
        q_text_run = q_text_para.add_run(f"{soru_metni.strip()}")
        q_text_run.font.size = Pt(11)

        # --- GÖRSEL EKLEME (Eğer varsa) ---
        if isinstance(item, dict) and 'image_stream' in item and item['image_stream']:
            try:
                doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_picture(item['image_stream'], width=Inches(4.0))
                doc.add_paragraph('\n')
            except Exception as img_err:
                print(f"Resim ekleme hatası: {img_err}")

        # CEVAP ALANI (Açık uçlu için daha büyük alan)
        answer_table = doc.add_table(rows=1, cols=1)
        answer_table.style = 'Table Grid'
        answer_cell = answer_table.cell(0, 0)
        
        # Açık uçlu sorular için daha büyük cevap alanı
        if is_acik_uclu:
            answer_cell.height = Inches(3.5)  # Daha büyük alan
        else:
            answer_cell.height = Inches(2.0)
        
        # Cevap başlığı
        answer_para = answer_cell.paragraphs[0]
        answer_para.text = "Cevap:"
        answer_para.runs[0].font.size = Pt(10)
        answer_para.runs[0].bold = True
        
        # Öğrenciler için boş satırlar (açık uçlu için daha fazla)
        satir_sayisi = 12 if is_acik_uclu else 8
        for _ in range(satir_sayisi):
            answer_cell.add_paragraph()
        
        doc.add_paragraph('\n')

    # --- 4. KISIM: HAZIRLAYAN BÖLÜMÜ ---
    
    doc.add_paragraph('\n\n')
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Öğretmen adını metadata'dan al
    ogretmen_adi = metadata.get('ogretmen', 'Öğretmen Adı')
    hazirlayan_run = footer_para.add_run(f"Hazırlayan:  {ogretmen_adi}")
    hazirlayan_run.font.size = Pt(11)
    hazirlayan_run.bold = True
    
    footer_para.add_run('\n')
    
    ogretmen_run = footer_para.add_run(f"({metadata.get('ders', 'Ders')} Öğretmeni)")
    ogretmen_run.font.size = Pt(11)

    # --- 5. KISIM: CEVAP ANAHTARI (Ayrı Sayfa) ---
    
    doc.add_page_break()
    
    answer_header = doc.add_paragraph()
    answer_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    answer_run = answer_header.add_run("CEVAP ANAHTARI (ÖĞRETMEN İÇİN)")
    answer_run.font.size = Pt(16)
    answer_run.bold = True
    answer_run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
    doc.add_paragraph('\n')
    
    # Cevapları listele (Açık uçlu formatında)
    for i, item in enumerate(questions[:10], 1):
        if isinstance(item, str):
            soru_metni = item
            beklenen_cozum = None
            puanlama_kriterleri = None
            detayli_cozum = None
            secenekler = []
            is_test = False
        else:
            soru_metni = item.get('soru_metni', item.get('soru', f"Soru {i} metni okunamadı."))
            beklenen_cozum = item.get('beklenen_cozum')
            puanlama_kriterleri = item.get('puanlama_kriterleri')
            detayli_cozum = item.get('detayli_cozum', item.get('detailed_solution'))
            secenekler = item.get('secenekler', [])
            is_test = secenekler and len(secenekler) > 0
        
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"{i}. Soru: {soru_metni}")
        q_run.bold = True
        q_run.font.size = Pt(11)
        
        if is_test:
            # TEST SORUSU FORMATI
            sec_para = doc.add_paragraph()
            sec_label = sec_para.add_run("Seçenekler:\n")
            sec_label.bold = True
            sec_label.font.size = Pt(10)
            for secenek in secenekler:
                s_run = sec_para.add_run(f"  {secenek}\n")
                s_run.font.size = Pt(10)
            
            cevap_metni = item.get('dogru_cevap', item.get('cevap', "Cevap bulunamadı"))
            a_para = doc.add_paragraph()
            a_run = a_para.add_run(f"Doğru Cevap: {cevap_metni}")
            a_run.font.size = Pt(11)
            a_run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
            a_run.bold = True
        else:
            # AÇIK UÇLU SORU FORMATI
            a_para = doc.add_paragraph()
            a_run = a_para.add_run("📋 Açık Uçlu Soru")
            a_run.font.size = Pt(10)
            a_run.bold = True
            a_run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)

        # Beklenen Çözüm (açık uçlu için)
        if beklenen_cozum:
            bc_para = doc.add_paragraph()
            bc_label = bc_para.add_run("✅ Beklenen Model Cevap:\n")
            bc_label.bold = True
            bc_label.font.size = Pt(10)
            bc_label.font.color.rgb = RGBColor(0x00, 0x90, 0x00)
            
            bc_text = bc_para.add_run(beklenen_cozum)
            bc_text.font.size = Pt(10)
            bc_text.italic = True

        # Puanlama Kriterleri (açık uçlu için)
        if puanlama_kriterleri:
            pk_para = doc.add_paragraph()
            pk_label = pk_para.add_run(f"📊 Puanlama Kriterleri:\n")
            pk_label.bold = True
            pk_label.font.size = Pt(10)
            pk_label.font.color.rgb = RGBColor(0x80, 0x40, 0x00)
            
            pk_text = pk_para.add_run(puanlama_kriterleri)
            pk_text.font.size = Pt(9)

        # Detaylı Çözüm (varsa - test için veya ek bilgi)
        if detayli_cozum and not beklenen_cozum:
            ds_para = doc.add_paragraph()
            ds_label = ds_para.add_run("📝 Detaylı Çözüm:\n")
            ds_label.bold = True
            ds_label.font.size = Pt(10)
            ds_label.font.color.rgb = RGBColor(0x00, 0x60, 0x90)
            
            ds_text = ds_para.add_run(detayli_cozum)
            ds_text.font.size = Pt(10)
            ds_text.italic = True
        
        # Bloom Değerlendirmesi (sadece cevap anahtarında)
        if isinstance(item, dict):
            bloom_level = item.get('bloom_seviyesi', item.get('bloom_level', ''))
            if bloom_level:
                bl_para = doc.add_paragraph()
                bl_run = bl_para.add_run(f"📚 Bloom Seviyesi: {bloom_level}")
                bl_run.font.size = Pt(9)
                bl_run.font.color.rgb = RGBColor(0x00, 0x60, 0x90)
                bl_run.italic = True
        
        # Not: İpucu ve yaygın hatalar cevap anahtarından kaldırıldı
        
        doc.add_paragraph()

    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    return docx_buffer
